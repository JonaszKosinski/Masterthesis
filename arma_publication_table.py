import warnings
import pandas as pd
from statsmodels.tsa.arima.model import ARIMA
from docx import Document

warnings.filterwarnings("ignore")

# ============================================================
# SETTINGS
# ============================================================
MASTER_PATH = "DATA/Processed/armax_master_dataframe.xlsx"
OUT_XLSX = "DATA/Processed/arma_publication_table.xlsx"
OUT_DOCX = "DATA/Processed/arma_publication_table.docx"

COUNTRIES = ["Thailand", "Philippines", "Korea", "Indonesia"]

MODEL_SPEC = {
    "Thailand": (1, 0, 0),      # AR(1)
    "Philippines": (1, 0, 1),   # ARMA(1,1)
    "Korea": (1, 0, 0),         # AR(1)
    "Indonesia": (1, 0, 0)      # AR(1)
}

# ============================================================
# HELPERS
# ============================================================
def load_master(path: str) -> pd.DataFrame:
    df = pd.read_excel(path, index_col=0)
    df.index = pd.to_datetime(df.index, errors="coerce")
    df = df.sort_index()
    return df


def get_inflation(master: pd.DataFrame, country: str) -> pd.Series:
    col = f"{country}_infl"
    if col not in master.columns:
        raise ValueError(f"Column not found: {col}")

    s = pd.to_numeric(master[col], errors="coerce").dropna().copy()
    s.index = pd.DatetimeIndex(s.index)
    s = s.asfreq("MS").dropna()

    if len(s) == 0:
        raise ValueError(f"{country}: empty series after processing.")

    return s


def fit_model(series: pd.Series, order: tuple):
    model = ARIMA(series, order=order, trend="c")
    result = model.fit(method_kwargs={"maxiter": 500})
    return result


def starify(coef, pval, decimals=4):
    if pd.isna(coef):
        return ""
    stars = ""
    if pval < 0.01:
        stars = "***"
    elif pval < 0.05:
        stars = "**"
    elif pval < 0.10:
        stars = "*"
    return f"{coef:.{decimals}f}{stars}"


def se_fmt(se, decimals=4):
    if pd.isna(se):
        return ""
    return f"({se:.{decimals}f})"


def build_country_results(master: pd.DataFrame):
    results = {}

    for country in COUNTRIES:
        series = get_inflation(master, country)
        order = MODEL_SPEC[country]
        res = fit_model(series, order)

        results[country] = {
            "nobs": int(res.nobs),
            "params": res.params.to_dict(),
            "bse": res.bse.to_dict(),
            "pvalues": res.pvalues.to_dict(),
            "aic": res.aic,
            "bic": res.bic
        }

    return results


def get_param_block(country_result: dict, param_name: str):
    coef = country_result["params"].get(param_name, float("nan"))
    se = country_result["bse"].get(param_name, float("nan"))
    pval = country_result["pvalues"].get(param_name, float("nan"))

    if pd.isna(coef):
        return "", ""
    return starify(coef, pval), se_fmt(se)


def build_publication_table(results: dict) -> pd.DataFrame:
    rows = []

    variables = [
        ("Constant", "const"),
        ("AR(1)", "ar.L1"),
        ("MA(1)", "ma.L1"),
        ("Sigma²", "sigma2"),
    ]

    for label, param in variables:
        coef_row = {"Variable": label}
        se_row = {"Variable": ""}

        for country in COUNTRIES:
            coef_text, se_text = get_param_block(results[country], param)
            coef_row[country] = coef_text
            se_row[country] = se_text

        rows.append(coef_row)
        rows.append(se_row)

    nobs_row = {"Variable": "Observations"}
    aic_row = {"Variable": "AIC"}
    bic_row = {"Variable": "BIC"}

    for country in COUNTRIES:
        nobs_row[country] = results[country]["nobs"]
        aic_row[country] = f'{results[country]["aic"]:.3f}'
        bic_row[country] = f'{results[country]["bic"]:.3f}'

    rows.append(nobs_row)
    rows.append(aic_row)
    rows.append(bic_row)

    return pd.DataFrame(rows)


def export_to_word(df: pd.DataFrame, filepath: str):
    doc = Document()
    doc.add_heading("Baseline ARMA Estimation Results", level=1)

    p = doc.add_paragraph()
    p.add_run("Notes. ").bold = True
    p.add_run("Entries report coefficient estimates with standard errors in parentheses. ")
    p.add_run("*** p<0.01, ** p<0.05, * p<0.10.")

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    # header
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    # body
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            cells[j].text = "" if pd.isna(val) else str(val)

    doc.save(filepath)


# ============================================================
# MAIN
# ============================================================
def main():
    master = load_master(MASTER_PATH)
    results = build_country_results(master)
    pub_table = build_publication_table(results)

    # Export to Excel
    with pd.ExcelWriter(OUT_XLSX, engine="openpyxl") as writer:
        pub_table.to_excel(writer, sheet_name="ARMA_Table", index=False)

    # Export to Word
    export_to_word(pub_table, OUT_DOCX)

    print(pub_table.to_string(index=False))
    print(f"\nSaved Excel table to: {OUT_XLSX}")
    print(f"Saved Word table to: {OUT_DOCX}")


if __name__ == "__main__":
    main()